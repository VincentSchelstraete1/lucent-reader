import io

import docx
import pytest

from app.ingestion import DocumentExtractionError, DocxDocumentIngestor
from app.ingestion.docx_ingestor import DOCX_MIME_TYPE
from app.main import app
from app.normalization import normalize_document
from app.routers.ingestion import get_docx_ingestor


def _tiny_png() -> bytes:
    import struct
    import zlib

    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body))

    signature = b"\x89PNG\r\n\x1a\n"
    header = struct.pack(">IIBBBBB", 2, 2, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * 2 for _ in range(2))
    data = zlib.compress(raw)
    return signature + chunk(b"IHDR", header) + chunk(b"IDAT", data) + chunk(b"IEND", b"")


def _sample_docx(*, include_image: bool = False) -> bytes:
    document = docx.Document()
    document.add_paragraph("Document Title", style="Title")
    document.add_paragraph("Cache Associativity", style="Heading 1")
    document.add_paragraph("Direct-mapped caches place a block in exactly one location.")
    document.add_paragraph("First bullet point", style="List Bullet")
    document.add_paragraph("Second bullet point", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Level"
    table.cell(0, 1).text = "Latency"
    table.cell(1, 0).text = "L1"
    table.cell(1, 1).text = "4 cycles"
    if include_image:
        document.add_picture(io.BytesIO(_tiny_png()), width=None)
    document.add_paragraph("Closing paragraph after the table.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class StubDocxIngestor:
    def __init__(self, *, error: bool = False) -> None:
        self.error = error
        self.calls: list[tuple[bytes, str]] = []

    def ingest_docx(self, docx_bytes: bytes, *, filename: str):
        self.calls.append((docx_bytes, filename))
        if self.error:
            raise DocumentExtractionError("docx extractor detail that must not leak")
        from app.ingestion import RawContentBlock, RawDocument, RawPage, SourceLocation

        block = RawContentBlock(
            id="docx-paragraph-1",
            page_number=None,
            type="text",
            bbox=None,
            reading_order=0,
            text="stub paragraph",
            location=SourceLocation(kind="document", sequence_id="paragraph-1"),
        )
        return RawDocument(
            source_type="docx",
            filename=filename,
            page_count=1,
            markdown="stub paragraph",
            pages=[RawPage(None, "stub paragraph", [block], location=SourceLocation(kind="document"))],
            images=[],
        )


@pytest.fixture()
def stub_docx_ingestor():
    stub = StubDocxIngestor()
    app.dependency_overrides[get_docx_ingestor] = lambda: stub
    yield stub
    app.dependency_overrides.pop(get_docx_ingestor, None)


def test_docx_reaches_normalized_document_with_headings_lists_tables():
    raw = DocxDocumentIngestor().ingest_docx(_sample_docx(), filename="notes.docx")
    assert raw.source_type == "docx"
    assert raw.page_count == 1
    assert len(raw.pages) == 1
    assert raw.pages[0].page_number is None

    normalized = normalize_document(raw)
    types = [block.type for block in normalized.pages[0].blocks]
    assert "heading" in types
    assert "list" in types
    assert "table" in types
    assert "paragraph" in types

    table_block = next(block for block in normalized.pages[0].blocks if block.type == "table")
    assert "L1" in (table_block.text or "")
    assert "4 cycles" in (table_block.text or "")


def test_docx_paragraph_and_table_provenance_uses_document_location_not_page_number():
    raw = DocxDocumentIngestor().ingest_docx(_sample_docx(), filename="notes.docx")
    for block in raw.pages[0].blocks:
        assert block.page_number is None
        assert block.location is not None
        assert block.location.kind == "document"
        assert block.location.index is None
        assert block.location.sequence_id is not None

    normalized = normalize_document(raw)
    for block in normalized.pages[0].blocks:
        assert block.source.page_start is None
        assert block.source.page_end is None
        assert all(location.kind == "document" for location in block.source.locations)


def test_docx_image_is_extracted_with_document_provenance():
    raw = DocxDocumentIngestor().ingest_docx(_sample_docx(include_image=True), filename="notes.docx")
    assert len(raw.images) == 1
    image = raw.images[0]
    assert image.page_number is None
    assert image.location is not None
    assert image.location.kind == "document"
    assert image.location.sequence_id is not None
    assert image.width == 2
    assert image.height == 2
    assert image.mime_type == "image/png"
    assert image.asset_reference.startswith("data:image/png;base64,")

    normalized = normalize_document(raw)
    assert len(normalized.images) == 1
    assert normalized.images[0].location is not None
    assert normalized.images[0].location.kind == "document"


def test_docx_normalization_event_with_no_page_number_serializes_safely():
    # A DOCX paragraph with an embedded manual line break has the same shape
    # that triggers hyphenation/line-wrap repair on PDF text - regression test
    # for NormalizationEvent.page_number being None (not int) for this format.
    document = docx.Document()
    document.add_paragraph(
        "This line wraps because of an artificial hyphen at the end, re-\nconfigurable systems are common."
    )
    buffer = io.BytesIO()
    document.save(buffer)

    raw = DocxDocumentIngestor().ingest_docx(buffer.getvalue(), filename="wrap.docx")
    normalized = normalize_document(raw)
    assert normalized.normalization_metadata.events
    assert normalized.normalization_metadata.events[0].page_number is None

    from app.schemas.ingestion import PdfIngestionResponse

    response = PdfIngestionResponse.from_documents(raw, normalized)
    assert response.normalized.normalization_metadata.events[0].page_number is None


def test_docx_extraction_failure_raises_document_extraction_error():
    with pytest.raises(DocumentExtractionError):
        DocxDocumentIngestor().ingest_docx(b"not a real docx file", filename="broken.docx")


def test_docx_upload_endpoint_returns_normalized_document(client, stub_docx_ingestor):
    response = client.post(
        "/ingestion/docx",
        files={"file": ("notes.docx", b"PK\x03\x04stub docx bytes", DOCX_MIME_TYPE)},
    )
    assert response.status_code == 200
    result = response.json()
    assert result["source_type"] == "docx"
    assert result["pages"][0]["page_number"] is None
    assert result["pages"][0]["location"]["kind"] == "document"
    assert result["pages"][0]["blocks"][0]["location"]["sequence_id"] == "paragraph-1"
    assert result["normalized"]["pages"][0]["page_number"] is None
    assert stub_docx_ingestor.calls == [(b"PK\x03\x04stub docx bytes", "notes.docx")]


def test_docx_upload_rejects_wrong_media_type(client, stub_docx_ingestor):
    response = client.post("/ingestion/docx", files={"file": ("notes.txt", b"not docx", "text/plain")})
    assert response.status_code == 415
    assert response.json()["detail"]["code"] == "unsupported_file_type"
    assert stub_docx_ingestor.calls == []


def test_docx_upload_rejects_invalid_zip_signature(client, stub_docx_ingestor):
    response = client.post(
        "/ingestion/docx",
        files={"file": ("notes.docx", b"not a zip file", DOCX_MIME_TYPE)},
    )
    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "invalid_docx"


def test_docx_upload_requires_authentication(unauthenticated_client, stub_docx_ingestor):
    response = unauthenticated_client.post(
        "/ingestion/docx",
        files={"file": ("notes.docx", b"PK\x03\x04stub docx bytes", DOCX_MIME_TYPE)},
    )
    assert response.status_code == 401
    assert stub_docx_ingestor.calls == []


def test_docx_extraction_failure_returns_safe_structured_error(client):
    app.dependency_overrides[get_docx_ingestor] = lambda: StubDocxIngestor(error=True)
    try:
        response = client.post(
            "/ingestion/docx",
            files={"file": ("broken.docx", b"PK\x03\x04stub docx bytes", DOCX_MIME_TYPE)},
        )
    finally:
        app.dependency_overrides.pop(get_docx_ingestor, None)
    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "docx_extraction_failed"
    assert "extractor detail" not in response.text


@pytest.mark.integration
def test_real_docx_extractor_returns_pages_markdown_and_provenance():
    result = DocxDocumentIngestor().ingest_docx(_sample_docx(), filename="fixture.docx")
    assert "Direct-mapped caches" in result.pages[0].text
    assert "Direct-mapped caches" in result.markdown
    assert result.page_count == 1
